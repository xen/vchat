from __future__ import annotations

import html
import logging
import re
import zipfile
from collections import Counter
from io import BytesIO
from pathlib import Path
from typing import Any
from urllib.parse import parse_qs, unquote, urlparse

from bs4 import BeautifulSoup
from docx import Document
from pypdf import PdfReader

from jobs.documents.types import guess_document_type

logger = logging.getLogger(__name__)

BOILERPLATE_TAGS = ("header", "footer", "nav", "aside")
BOILERPLATE_HINTS = (
    "cookie",
    "cookies",
    "banner",
    "menu",
    "navbar",
    "nav",
    "header",
    "footer",
    "sidebar",
    "breadcrumbs",
    "breadcrumb",
    "advert",
    "promo",
    "share",
    "social",
)
BOILERPLATE_ATTR_HINTS = BOILERPLATE_HINTS + (
    "modal",
    "popup",
    "overlay",
    "dialog",
    "search",
    "login",
    "signin",
    "signup",
    "register",
    "auth",
)
AUTH_HEADING_HINTS = {
    "авторизация",
    "войти",
    "вход",
    "зарегистрироваться",
    "регистрация",
    "восстановить пароль",
    "забыли пароль",
    "login",
    "sign in",
    "sign up",
    "register",
    "forgot password",
}
CODE_FENCE_RE = re.compile(r"^```(?P<lang>[a-zA-Z0-9_+-]*)\s*$")
WORD_DOCUMENT_TYPES = {
    "application/msword",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}
PDF_DOCUMENT_TYPES = {"application/pdf"}
DOCX_DOCUMENT_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}
LEGACY_DOC_DOCUMENT_TYPES = {"application/msword"}
BROAD_BINARY_DOCUMENT_TYPES = {
    "",
    "application/octet-stream",
    "binary/octet-stream",
    "application/binary",
    "application/download",
    "application/force-download",
    "application/x-download",
}


def normalize_markdown(text: str) -> str:
    lines = text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    normalized: list[str] = []
    blank_count = 0
    in_code_block = False
    for raw_line in lines:
        line = raw_line.rstrip()
        if CODE_FENCE_RE.match(line.strip()):
            in_code_block = not in_code_block
            blank_count = 0
            normalized.append(line.strip() or "```")
            continue
        if in_code_block:
            normalized.append(line)
            continue
        if not line.strip():
            blank_count += 1
            if blank_count <= 1:
                normalized.append("")
            continue
        blank_count = 0
        normalized.append(line)
    return "\n".join(normalized).strip()


def _looks_like_table(lines: list[str], index: int) -> bool:
    if index + 1 >= len(lines):
        return False
    line = lines[index].strip()
    separator = lines[index + 1].strip()
    if "|" not in line or "|" not in separator:
        return False
    cells = [cell.strip() for cell in separator.strip("|").split("|")]
    if not cells:
        return False
    valid = [cell for cell in cells if re.fullmatch(r":?-{3,}:?", cell)]
    return len(valid) >= max(1, len(cells) - 1)


def build_outline(structure: list[dict[str, Any]]) -> list[dict[str, Any]]:
    outline: list[dict[str, Any]] = []
    for block in structure:
        if block.get("type") != "heading":
            continue
        outline.append(
            {
                "level": block.get("level", 1),
                "content": block.get("content", ""),
                "section_path": block.get("section_path"),
            }
        )
    return outline


def build_structure_from_markdown(
    markdown: str,
) -> tuple[list[dict[str, Any]], dict[str, Any]]:
    lines = markdown.splitlines()
    blocks: list[dict[str, Any]] = []
    heading_stack: list[str] = []
    paragraph_lines: list[str] = []
    table_count = 0
    code_count = 0
    list_count = 0
    paragraph_count = 0
    heading_count = 0
    code_lang = ""
    code_lines: list[str] = []
    ordered_items: list[str] = []
    unordered_items: list[str] = []

    def current_section_path() -> str | None:
        if not heading_stack:
            return None
        return " / ".join(heading_stack)

    def flush_paragraph() -> None:
        nonlocal paragraph_lines, paragraph_count
        text = "\n".join(paragraph_lines).strip()
        paragraph_lines = []
        if not text:
            return
        blocks.append(
            {
                "type": "paragraph",
                "content": text,
                "section_path": current_section_path(),
            }
        )
        paragraph_count += 1

    def flush_list() -> None:
        nonlocal ordered_items, unordered_items, list_count
        if ordered_items:
            blocks.append(
                {
                    "type": "list",
                    "ordered": True,
                    "items": ordered_items[:],
                    "section_path": current_section_path(),
                }
            )
            ordered_items = []
            list_count += 1
        if unordered_items:
            blocks.append(
                {
                    "type": "list",
                    "ordered": False,
                    "items": unordered_items[:],
                    "section_path": current_section_path(),
                }
            )
            unordered_items = []
            list_count += 1

    index = 0
    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        fence = CODE_FENCE_RE.match(stripped)
        if fence:
            flush_paragraph()
            flush_list()
            if code_lines:
                blocks.append(
                    {
                        "type": "code",
                        "language": code_lang or "",
                        "content": "\n".join(code_lines).strip(),
                        "section_path": current_section_path(),
                    }
                )
                code_count += 1
                code_lines = []
                code_lang = ""
            else:
                code_lang = fence.group("lang") or ""
            index += 1
            while index < len(lines):
                next_line = lines[index]
                if CODE_FENCE_RE.match(next_line.strip()):
                    break
                code_lines.append(next_line.rstrip())
                index += 1
            if code_lines:
                blocks.append(
                    {
                        "type": "code",
                        "language": code_lang or "",
                        "content": "\n".join(code_lines).strip(),
                        "section_path": current_section_path(),
                    }
                )
                code_count += 1
                code_lines = []
                code_lang = ""
            if index < len(lines) and CODE_FENCE_RE.match(lines[index].strip()):
                index += 1
            continue

        heading = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading:
            flush_paragraph()
            flush_list()
            level = len(heading.group(1))
            title = heading.group(2).strip()
            heading_stack = heading_stack[: level - 1]
            heading_stack.append(title)
            blocks.append(
                {
                    "type": "heading",
                    "level": level,
                    "content": title,
                    "section_path": " / ".join(heading_stack),
                }
            )
            heading_count += 1
            index += 1
            continue

        if _looks_like_table(lines, index):
            flush_paragraph()
            flush_list()
            table_lines = [lines[index].rstrip(), lines[index + 1].rstrip()]
            index += 2
            while index < len(lines) and "|" in lines[index]:
                table_lines.append(lines[index].rstrip())
                index += 1
            blocks.append(
                {
                    "type": "table",
                    "content": "\n".join(table_lines).strip(),
                    "caption": None,
                    "section_path": current_section_path(),
                }
            )
            table_count += 1
            continue

        ordered = re.match(r"^\d+\.\s+(.*)$", stripped)
        unordered = re.match(r"^[-*+]\s+(.*)$", stripped)
        if ordered:
            flush_paragraph()
            unordered_items = []
            ordered_items.append(ordered.group(1).strip())
            index += 1
            continue
        if unordered:
            flush_paragraph()
            ordered_items = []
            unordered_items.append(unordered.group(1).strip())
            index += 1
            continue

        if not stripped:
            flush_paragraph()
            flush_list()
            index += 1
            continue

        paragraph_lines.append(line.rstrip())
        index += 1

    flush_paragraph()
    flush_list()
    metadata = {
        "word_count": len(markdown.split()),
        "table_count": table_count,
        "heading_count": heading_count,
        "list_count": list_count,
        "paragraph_count": paragraph_count,
        "code_block_count": code_count,
    }
    return blocks, metadata


def _coerce_title(candidate: str | None, structure: list[dict[str, Any]]) -> str | None:
    normalized_candidate = normalize_title_candidate(candidate)
    headings: list[str] = []
    for block in structure:
        if block.get("type") != "heading":
            continue
        value = normalize_title_candidate(block.get("content"))
        if value:
            headings.append(value)

    if normalized_candidate:
        if any(
            heading.casefold() == normalized_candidate.casefold()
            for heading in headings
        ):
            return normalized_candidate
        if headings and headings[0].casefold() in AUTH_HEADING_HINTS:
            return normalized_candidate

    for value in headings:
        return value
    return normalized_candidate


def normalize_title_candidate(candidate: str | None) -> str | None:
    if not candidate:
        return None

    cleaned = html.unescape(str(candidate))
    cleaned = re.sub(r"\s+", " ", cleaned).strip()
    if not cleaned:
        return None

    lowered = cleaned.lower()
    bad_markers = (
        "fatal error",
        "traceback (most recent call last)",
        "exception:",
        "warning:",
        "<script",
    )
    if any(marker in lowered for marker in bad_markers):
        return None

    # Guard against malformed HTML where <title> may absorb page body.
    if len(cleaned) > 220 or len(cleaned.split()) > 30:
        return None

    # Reject URL slugs: all-lowercase ASCII with no spaces (e.g. Docling's document.name).
    # Mixed-case strings like "Python" or "HTML5" are legitimate titles and pass through.
    if len(cleaned) >= 4 and re.fullmatch(r"[a-z][a-z0-9\-_]*", cleaned):
        return None

    return cleaned[:512]


def normalize_file_metadata_title_candidate(candidate: str | None) -> str | None:
    if not candidate:
        return None

    cleaned = html.unescape(str(candidate))
    cleaned = re.sub(r"\s+", " ", cleaned).strip()
    if not cleaned:
        return None

    lowered = cleaned.lower()
    bad_markers = (
        "fatal error",
        "traceback (most recent call last)",
        "exception:",
        "warning:",
        "<script",
    )
    if any(marker in lowered for marker in bad_markers):
        return None
    if len(cleaned) > 220 or len(cleaned.split()) > 30:
        return None

    return cleaned[:512]


def build_document_payload(
    *,
    content: str,
    title: str | None,
    extractor: str,
    fallback_used: bool,
    degraded_mode: bool,
    content_type: str | None = None,
    doc_type: str | None = None,
    extra_meta: dict[str, Any] | None = None,
) -> tuple[str, str | None, dict[str, Any]]:
    normalized = normalize_markdown(content)
    structure, structure_meta = build_structure_from_markdown(normalized)
    outline = build_outline(structure)
    meta: dict[str, Any] = {
        "structure": structure,
        "outline": outline,
        "extraction": {
            "extractor": extractor,
            "fallback_used": fallback_used,
            "degraded_mode": degraded_mode,
            "table_count": structure_meta["table_count"],
            "word_count": structure_meta["word_count"],
            "boilerplate_removed_count": 0,
            "reason": "plain_text_fallback" if degraded_mode else None,
        },
    }
    if content_type:
        meta["content_type"] = content_type
    if doc_type:
        meta["doc_type"] = doc_type
    if extra_meta:
        meta.update(extra_meta)
    normalized_title = _coerce_title(title, structure)
    return normalized, normalized_title, meta


def _title_from_filename(filename: str | None) -> str | None:
    if not filename:
        return None
    stem = Path(unquote(filename)).stem.replace("_", " ").replace("-", " ").strip()
    if not stem:
        return None
    return normalize_title_candidate(stem) or re.sub(r"\s+", " ", stem).strip()[:512]


def _title_from_url(source_url: str) -> str | None:
    parsed = urlparse(source_url)
    filename = Path(unquote(parsed.path)).name
    if "." not in filename:
        params = parse_qs(parsed.query)
        for key in ("filename", "file", "name"):
            values = params.get(key) or []
            for value in values:
                candidate = Path(unquote(value)).name
                if candidate:
                    filename = candidate
                    break
            if "." in filename:
                break
    if not filename:
        return None
    return _title_from_filename(filename)


def _extract_pdf_text(reader: PdfReader) -> tuple[str, dict[str, Any]]:
    pages: list[str] = []
    for index, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if not text:
            continue
        pages.append(f"## Page {index}\n\n{text}")
    return "\n\n".join(pages), {"page_count": len(reader.pages)}


def _extract_pdf_metadata_title(reader: PdfReader) -> str | None:
    metadata = reader.metadata
    if metadata is None:
        return None
    return normalize_file_metadata_title_candidate(getattr(metadata, "title", None))


def _extract_docx_metadata_title(document: Document) -> str | None:
    return normalize_file_metadata_title_candidate(document.core_properties.title)


def _extract_docx_text(document: Document) -> tuple[str, dict[str, Any]]:
    parts: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in document.tables:
        rows: list[list[str]] = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                rows.append(cells)
        if not rows:
            continue
        width = max(len(row) for row in rows)
        normalized_rows = [row + [""] * (width - len(row)) for row in rows]
        header = normalized_rows[0]
        parts.append(f"| {' | '.join(header)} |")
        parts.append(f"| {' | '.join(['---'] * width)} |")
        parts.extend(f"| {' | '.join(row)} |" for row in normalized_rows[1:])

    return "\n\n".join(parts), {
        "paragraph_count": len(document.paragraphs),
        "table_count": len(document.tables),
    }


def _normalized_content_type(content_type: str | None) -> str:
    return (content_type or "").split(";", 1)[0].strip().lower()


def _sniff_binary_document_kind(raw_body: bytes) -> str | None:
    if raw_body.startswith(b"%PDF-"):
        return "pdf"
    if raw_body.startswith(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"):
        return "legacy_doc"
    if zipfile.is_zipfile(BytesIO(raw_body)):
        with zipfile.ZipFile(BytesIO(raw_body)) as archive:
            names = set(archive.namelist())
        if "[Content_Types].xml" in names and "word/document.xml" in names:
            return "docx"
    return None


def detect_binary_document_kind(
    raw_body: bytes,
    content_type: str | None,
) -> str | None:
    normalized_content_type = _normalized_content_type(content_type)
    if normalized_content_type in PDF_DOCUMENT_TYPES:
        return "pdf"
    if normalized_content_type in DOCX_DOCUMENT_TYPES:
        return "docx"
    if normalized_content_type in LEGACY_DOC_DOCUMENT_TYPES:
        return "legacy_doc"
    if normalized_content_type in BROAD_BINARY_DOCUMENT_TYPES:
        return _sniff_binary_document_kind(raw_body)
    return None


def extract_binary_url_document(
    source_url: str,
    raw_body: bytes,
    content_type: str | None = None,
) -> tuple[str, str | None, dict[str, Any]]:
    document_kind = detect_binary_document_kind(raw_body, content_type)
    doc_type = "office" if document_kind else guess_document_type(None, content_type)
    title = _title_from_url(source_url)

    if document_kind == "pdf":
        reader = PdfReader(BytesIO(raw_body))
        content, file_meta = _extract_pdf_text(reader)
        title = _extract_pdf_metadata_title(reader) or title
        extractor = "pypdf"
        fallback_used = False
    elif document_kind == "docx":
        document = Document(BytesIO(raw_body))
        title = _extract_docx_metadata_title(document) or title
        content, file_meta = _extract_docx_text(document)
        extractor = "python-docx"
        fallback_used = False
    elif document_kind == "legacy_doc":
        raise ValueError("Legacy .doc files are not supported by the current extractor")
    else:
        raise ValueError(
            f"Unsupported downloadable document type: {content_type or source_url}"
        )

    content, _document_title, meta = build_document_payload(
        content=content,
        title=title,
        extractor=extractor,
        fallback_used=fallback_used,
        degraded_mode=False,
        content_type=content_type,
        doc_type=doc_type,
        extra_meta={"file": file_meta},
    )
    return content, title, meta


def _strip_boilerplate(soup: BeautifulSoup) -> int:
    removed = 0
    for tag_name in BOILERPLATE_TAGS:
        for node in soup.find_all(tag_name):
            node.decompose()
            removed += 1
    for node in soup.find_all(["dialog", "template"]):
        node.decompose()
        removed += 1
    for form in list(soup.find_all("form")):
        if form.find("input", attrs={"type": re.compile(r"password", re.I)}):
            form.decompose()
            removed += 1
    for node in list(soup.find_all(True)):
        attrs_dict = getattr(node, "attrs", None) or {}
        attrs = " ".join(
            [
                " ".join(attrs_dict.get("class", [])),
                attrs_dict.get("id", "") or "",
                attrs_dict.get("role", "") or "",
                attrs_dict.get("aria-label", "") or "",
            ]
        ).lower()
        if not attrs:
            continue
        if any(hint in attrs for hint in BOILERPLATE_ATTR_HINTS):
            node.decompose()
            removed += 1
    return removed


def _html_to_markdown_like(soup: BeautifulSoup) -> tuple[str, int]:
    boilerplate_removed = _strip_boilerplate(soup)
    container = soup.body or soup
    lines: list[str] = []
    for element in container.find_all(
        [
            "h1",
            "h2",
            "h3",
            "h4",
            "h5",
            "h6",
            "p",
            "li",
            "pre",
            "code",
            "table",
        ]
    ):
        text = element.get_text("\n", strip=True)
        if not text:
            continue
        if re.fullmatch(r"h[1-6]", element.name):
            level = int(element.name[1])
            lines.append(f"{'#' * level} {text}")
            lines.append("")
            continue
        if element.name == "li":
            prefix = "- "
            parent = element.parent.name if element.parent else ""
            if parent == "ol":
                siblings = [
                    sibling
                    for sibling in element.parent.find_all("li", recursive=False)
                ]
                index = siblings.index(element) + 1 if element in siblings else 1
                prefix = f"{index}. "
            lines.append(f"{prefix}{text}")
            continue
        if element.name == "table":
            rows = []
            for row in element.find_all("tr"):
                cells = row.find_all(["th", "td"])
                if not cells:
                    continue
                rows.append([cell.get_text(" ", strip=True) for cell in cells])
            if rows:
                header = rows[0]
                rows_md = [
                    f"| {' | '.join(header)} |",
                    f"| {' | '.join(['---'] * len(header))} |",
                ]
                rows_md.extend(f"| {' | '.join(row)} |" for row in rows[1:])
                lines.extend(rows_md)
                lines.append("")
            continue
        if element.name in {"pre", "code"}:
            lines.append("```")
            lines.append(text)
            lines.append("```")
            lines.append("")
            continue
        lines.append(text)
        lines.append("")
    markdown = normalize_markdown("\n".join(lines))
    return markdown, boilerplate_removed


def _extract_nav_title(soup: BeautifulSoup, url: str) -> str | None:
    """
    Find a self-referencing navigation link on the page and return its text.

    Many sites (especially SPAs) share a generic <title> across all pages
    but expose human-readable page names in navigation anchors that link back
    to the current URL path.  The first such anchor whose text passes
    normalize_title_candidate is used as the page-specific title.
    """
    from urllib.parse import urlparse

    parsed = urlparse(url)
    path = parsed.path.rstrip("/")
    if not path:
        return None

    for a in soup.find_all("a", href=True):
        href = (a.get("href") or "").rstrip("/")
        if href == path or href == url.rstrip("/"):
            text = a.get_text(" ", strip=True)
            candidate = normalize_title_candidate(text)
            if candidate:
                return candidate
    return None


def extract_url_document(
    source_url: str,
    html_body: str,
    content_type: str | None = None,
) -> tuple[str, str | None, dict[str, Any]]:
    doc_type = guess_document_type(source_url, "text/html")

    body = html_body
    soup = BeautifulSoup(body, "html.parser")
    html_title = soup.title.get_text(" ", strip=True)[:512] if soup.title else None

    # Prefer a self-referencing nav link over the generic site-wide <title>.
    nav_title = _extract_nav_title(soup, source_url)
    best_title = nav_title or html_title

    markdown, removed = _html_to_markdown_like(soup)
    if markdown.strip():
        normalized, normalized_title, meta = build_document_payload(
            content=markdown,
            title=best_title,
            extractor="beautifulsoup",
            fallback_used=False,
            degraded_mode=False,
            content_type=content_type,
            doc_type=doc_type,
        )
        meta["extraction"]["boilerplate_removed_count"] = removed
        return normalized, normalized_title, meta

    return build_document_payload(
        content=body,
        title=best_title,
        extractor="plain_text_html",
        fallback_used=True,
        degraded_mode=True,
        content_type=content_type,
        doc_type=doc_type,
    )


def summarize_structure(structure: list[dict[str, Any]]) -> dict[str, int]:
    counts = Counter(block.get("type", "unknown") for block in structure)
    return dict(counts)
